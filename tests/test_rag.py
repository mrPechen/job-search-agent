import pytest
from sqlalchemy import text
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine

from src.database.db_settings import Base
from src.database.models import User
from src.rag.ingest import build_embeddings, chunk_by_sections, extract_cv_text
from src.rag.retrieve import retrieve_relevant, store_chunks


def test_build_embeddings_ollama():
    """Ollama-эмбеддинги создаются без сетевого вызова."""
    emb = build_embeddings("ollama", "nomic-embed-text", "", "http://localhost:11434")
    assert emb.model == "nomic-embed-text"


def test_build_embeddings_unknown_provider_raises():
    with pytest.raises(ValueError):
        build_embeddings("unknown", "m", "", "")


def test_chunk_by_sections_splits_on_blank_lines():
    text = "Навыки: Python\nSQL\n\nОпыт:\nИнженер\n\nЦель: разработчик"
    chunks = chunk_by_sections(text)
    assert chunks == [
        "Навыки: Python\nSQL",
        "Опыт:\nИнженер",
        "Цель: разработчик",
    ]


def test_chunk_by_sections_splits_long_chunk():
    long_text = "word " * 400  # 2000 символов — больше max_chars (1500)
    chunks = chunk_by_sections(long_text)
    assert len(chunks) > 1
    assert all(len(chunk) <= 1500 for chunk in chunks)


async def test_extract_cv_text_docx(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_paragraph("Иван Иванов")
    doc.add_paragraph("Python-разработчик")
    path = tmp_path / "cv.docx"
    doc.save(str(path))

    text = await extract_cv_text(path)

    assert "Иван Иванов" in text
    assert "Python-разработчик" in text


@pytest.fixture(scope="module")
def pg_url():
    """Поднять контейнер pgvector или пропустить тест, если Docker недоступен."""
    try:
        from testcontainers.postgres import PostgresContainer

        container = PostgresContainer("pgvector/pgvector:pg16")
        container.start()
    except Exception as exc:  # noqa: BLE001 - любая ошибка запуска = нет Docker
        pytest.skip(f"Docker недоступен: {exc}")
    url = container.get_connection_url(driver="asyncpg")
    yield url
    container.stop()


async def test_store_and_retrieve_relevant(pg_url):
    engine = create_async_engine(pg_url)

    # Подготовка схемы: расширение vector + все таблицы Base
    async with engine.begin() as conn:
        await conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
        await conn.run_sync(Base.metadata.create_all)

    session_factory = async_sessionmaker(engine, expire_on_commit=False)

    # Документ ссылается на users.id, поэтому нужна запись пользователя
    async with session_factory() as session:
        session.add(User(id=1, telegram_id="t1"))
        await session.commit()

    # Чанк A близок к запросу, чанк B — противоположен
    emb_a = [1.0] * 768
    emb_b = [-1.0] * 768
    async with session_factory() as session:
        count = await store_chunks(
            session, 1, ["Python developer", "Sales manager"], [emb_a, emb_b]
        )
        await session.commit()
        assert count == 2

    async with session_factory() as session:
        result = await retrieve_relevant(session, 1, [1.0] * 768, top_k=2)
        assert len(result) == 2
        assert result[0].chunk_text == "Python developer"

    await engine.dispose()


async def test_profile_repository_upsert(pg_url):
    """Сохранение профиля: вторая запись обновляет существующую, не создавая дубль."""
    from sqlalchemy import select

    from src.database.models import Profile
    from src.rag.profile_repo import ProfileRepository
    from src.rag.schemas import ProfileData

    engine = create_async_engine(pg_url)
    async with engine.begin() as conn:
        await conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
        await conn.run_sync(Base.metadata.create_all)

    session_factory = async_sessionmaker(engine, expire_on_commit=False)

    async with session_factory() as session:
        session.add(User(id=2, telegram_id="t2"))
        await session.commit()

    repo = ProfileRepository(session_factory)
    await repo.save(2, ProfileData(skills=["python"], desired_role="dev"))
    await repo.save(2, ProfileData(skills=["python", "sql"], desired_role="dev"))

    async with session_factory() as session:
        result = await session.execute(select(Profile).where(Profile.user_id == 2))
        rows = result.scalars().all()

    assert len(rows) == 1
    assert rows[0].skills == ["python", "sql"]

    await engine.dispose()


async def test_resume_retriever_returns_chunk_texts(pg_url):
    """Ретривер эмбеддит запрос и возвращает тексты ближайших чанков."""
    from src.rag.retrieve import ResumeRetriever

    engine = create_async_engine(pg_url)
    async with engine.begin() as conn:
        await conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
        await conn.run_sync(Base.metadata.create_all)

    session_factory = async_sessionmaker(engine, expire_on_commit=False)

    async with session_factory() as session:
        session.add(User(id=3, telegram_id="t3"))
        await session.commit()

    async with session_factory() as session:
        await store_chunks(
            session, 3, ["Python developer", "Sales manager"], [[1.0] * 768, [-1.0] * 768]
        )
        await session.commit()

    class FakeEmbedder:
        async def aembed_query(self, query: str) -> list[float]:
            return [1.0] * 768

    retriever = ResumeRetriever(session_factory, FakeEmbedder())
    texts = await retriever(3, "python backend")

    assert texts[0] == "Python developer"

    await engine.dispose()
