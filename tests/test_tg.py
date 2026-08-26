from langgraph.checkpoint.memory import MemorySaver

from src.agent.capabilities import CapabilityChoice
from src.agent.graph import CandidateScore, build_graph
from src.rag.schemas import ProfileData
from src.tg.service import TgService


class FakeModel:
    """Поддельная текстовая модель: фиксированный ответ на любое сообщение."""

    async def ainvoke(self, messages):
        return type("Msg", (), {"content": "привет, чем помочь?"})()


class FakeGateway:
    """Поддельный gateway: возвращает предзаданные структурированные ответы."""

    def __init__(self, capability_key: str | None = "chat", score: float = 0.1) -> None:
        self.capability_key = capability_key
        self.score = score
        self.text_model = FakeModel()

    async def invoke_structured(self, model, messages, schema):
        name = schema.__name__
        if name == "CapabilityChoice":
            return CapabilityChoice(key=self.capability_key)
        if name == "CandidateScore":
            return CandidateScore(score=self.score, reason="ok")
        if name == "ProfileData":
            return ProfileData(skills=["Python"], desired_role="разработчик")
        raise ValueError(f"unexpected schema {name}")


class FakeSearcher:
    """Поддельный поиск: возвращает заранее заданный список вакансий."""

    def __init__(self, candidates: list[dict]) -> None:
        self._candidates = candidates

    async def __call__(self, user_id: int, query: str = "") -> list[dict]:
        return self._candidates


class _FakeResult:
    """Поддельный результат запроса: пользователь не найден."""

    def scalar_one_or_none(self):
        return None


class _FakeSession:
    """Поддельная асинхронная сессия: все операции — no-op, фиксирует add."""

    def __init__(self) -> None:
        self.added = []
        self._next_id = 1

    async def execute(self, *args, **kwargs):
        return _FakeResult()

    def add(self, obj) -> None:
        self.added.append(obj)

    async def commit(self) -> None:
        pass

    async def refresh(self, obj) -> None:
        if getattr(obj, "id", None) is None:
            obj.id = self._next_id
            self._next_id += 1

    async def __aenter__(self) -> "_FakeSession":
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb) -> None:
        pass


def _fake_session_factory():
    """Фабрика поддельных сессий: всегда создаёт нового пользователя."""

    def factory():
        return _FakeSession()

    return factory


class FakeSitesRepo:
    """Поддельный репозиторий сайтов: заданный список + фиксация добавлений."""

    def __init__(self, domains):
        self._domains = list(domains)
        self.added = []

    async def get_domains(self, user_id):
        return self._domains

    async def add_domains(self, user_id, domains):
        self.added.extend(domains)
        self._domains.extend(domains)


class FakeProfileRepo:
    """Поддельный репозиторий профиля: фиксирует вызовы save."""

    def __init__(self) -> None:
        self.saved: list[tuple[int, ProfileData]] = []

    async def save(self, user_id: int, data: ProfileData) -> None:
        self.saved.append((user_id, data))


def _build_service(
    gateway: FakeGateway,
    searcher: FakeSearcher | None = None,
    sites_repo: FakeSitesRepo | None = None,
    profile_repo: FakeProfileRepo | None = None,
) -> TgService:
    """Собрать TgService с фейками и in-memory чекпоинтером."""
    searcher = searcher or FakeSearcher([])
    sites_repo = sites_repo if sites_repo is not None else FakeSitesRepo(["hh.ru"])
    graph = build_graph(gateway, searcher, checkpointer=MemorySaver())
    return TgService(
        graph,
        _fake_session_factory(),
        gateway=gateway,
        sites_repo=sites_repo,
        profile_repo=profile_repo,
    )


async def test_chat_message_returns_reply():
    """Свободное сообщение: способность chat, возвращает текст ответа."""
    service = _build_service(FakeGateway(capability_key="chat"))
    reply = await service.handle_message("123", "привет")
    assert reply == "привет, чем помочь?"


async def test_search_message_triggers_hitl():
    """Высокий скоринг: первый вызов возвращает запрос на подтверждение отклика."""
    service = _build_service(
        FakeGateway(capability_key="search_job", score=0.9),
        FakeSearcher([{"title": "Python Dev"}]),
    )
    reply = await service.handle_message("123", "найди работу")
    assert "Подтвердить отклик" in reply
    assert "Python Dev" in reply


def test_is_approval_rejects_negation():
    """Сообщения с отрицанием («не отправляй») не считаются подтверждением."""
    service = _build_service(FakeGateway(capability_key="chat"))
    assert service._is_approval("да") is True
    assert service._is_approval("не отправляй") is False
    assert service._is_approval("нет") is False


def test_is_approval_does_not_match_inside_words():
    """Короткие слова («да», «ок») не должны совпадать внутри других слов."""
    service = _build_service(FakeGateway(capability_key="chat"))
    assert service._is_approval("подарок") is False
    assert service._is_approval("свидание") is False
    assert service._is_approval("блок") is False
    assert service._is_approval("интернет") is False


async def test_confirm_resumes_and_reports():
    """Подтверждение «да» возобновляет приостановленный граф и даёт итоговый отчёт."""
    service = _build_service(
        FakeGateway(capability_key="search_job", score=0.9),
        FakeSearcher([{"title": "Python Dev"}]),
    )
    await service.handle_message("123", "найди работу")
    reply = await service.handle_message("123", "да")
    assert reply == "Откликнулся на 1 вакансий"


async def test_onboarding_asks_for_sites_when_none():
    """Нет сайтов: бот спрашивает, на каких сайтах искать, и не запускает диспетчер."""
    service = _build_service(
        FakeGateway(capability_key="chat"), sites_repo=FakeSitesRepo([])
    )
    reply = await service.handle_message("123", "привет")
    assert "На каких сайтах" in reply


async def test_onboarding_saves_sites_from_message():
    """Сообщение с доменами при отсутствии сайтов сохраняет их и подтверждает."""
    sites_repo = FakeSitesRepo([])
    service = _build_service(FakeGateway(capability_key="chat"), sites_repo=sites_repo)
    reply = await service.handle_message("123", "ищи на hh.ru и habr.com")
    assert "hh.ru" in reply
    assert "habr.com" in reply
    assert "hh.ru" in sites_repo.added
    assert "habr.com" in sites_repo.added


async def test_unknown_request_replies_not_available():
    """Запрос без подходящей способности → честный ответ «не умею»."""
    service = _build_service(FakeGateway(capability_key=None))
    reply = await service.handle_message("123", "отправь мне письмо на почту")
    assert "Такого функционала пока нет" in reply


async def test_upload_resume_prompts_file():
    """Способность upload_resume просит прислать файл."""
    service = _build_service(FakeGateway(capability_key="upload_resume"))
    reply = await service.handle_message("123", "куда прислать резюме")
    assert "PDF или DOCX" in reply


async def test_handle_document_saves_profile(tmp_path):
    """Присланный DOCX парсится и сохраняется в репозиторий профиля."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Иван Иванов")
    doc.add_paragraph("Python-разработчик")
    path = tmp_path / "cv.docx"
    doc.save(str(path))

    profile_repo = FakeProfileRepo()
    service = _build_service(
        FakeGateway(capability_key="chat"), profile_repo=profile_repo
    )
    reply = await service.handle_document("123", "cv.docx", path.read_bytes())

    assert "Записал профиль" in reply
    assert len(profile_repo.saved) == 1
    user_id, data = profile_repo.saved[0]
    assert user_id == 1
    assert data.skills == ["Python"]


async def test_handle_document_rejects_bad_format(tmp_path):
    """Неподдерживаемый формат отклоняется без парсинга."""
    service = _build_service(FakeGateway(capability_key="chat"))
    reply = await service.handle_document("123", "cv.txt", b"hello")
    assert "PDF или DOCX" in reply
